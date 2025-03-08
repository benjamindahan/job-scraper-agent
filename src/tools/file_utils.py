"""
Utilities for extracting text from CV files (PDF, DOCX, etc.)
"""

import os
import logging
from typing import Optional, Dict, Any
import io

# Configure logging
logging.basicConfig(level=logging.INFO,
                   format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

def extract_text_from_file(file_path: Optional[str] = None,
                          file_content: Optional[bytes] = None,
                          file_name: Optional[str] = None) -> Dict[str, Any]:
    """
    Extract text from a CV file (PDF, DOCX, or plain text).

    Args:
        file_path: Path to the file (optional)
        file_content: Binary content of the file (optional)
        file_name: Name of the file, including extension (optional, used to determine file type)

    Returns:
        Dictionary with:
        - 'text': Extracted text as a string
        - 'success': Boolean indicating whether extraction was successful
        - 'error': Error message if extraction failed
    """
    if not file_path and not file_content:
        return {
            'text': '',
            'success': False,
            'error': 'No file path or content provided'
        }

    try:
        # Determine file extension
        if file_path:
            _, file_ext = os.path.splitext(file_path)
        elif file_name:
            _, file_ext = os.path.splitext(file_name)
        else:
            return {
                'text': '',
                'success': False,
                'error': 'Unable to determine file type'
            }

        file_ext = file_ext.lower()

        # Process based on file type
        if file_ext == '.pdf':
            try:
                from pypdf import PdfReader

                if file_path:
                    reader = PdfReader(file_path)
                else:
                    reader = PdfReader(io.BytesIO(file_content))

                text = ""
                for page in reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                return {
                    'text': text,
                    'success': True,
                    'error': None
                }
            except Exception as e:
                logger.error(f"Error extracting text from PDF: {str(e)}")
                return {
                    'text': '',
                    'success': False,
                    'error': f"PDF extraction error: {str(e)}"
                }

        elif file_ext == '.docx':
            try:
                import docx

                if file_path:
                    doc = docx.Document(file_path)
                else:
                    doc = docx.Document(io.BytesIO(file_content))

                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"

                return {
                    'text': text,
                    'success': True,
                    'error': None
                }
            except Exception as e:
                logger.error(f"Error extracting text from DOCX: {str(e)}")
                return {
                    'text': '',
                    'success': False,
                    'error': f"DOCX extraction error: {str(e)}"
                }

        else:
            # Handle text files or other formats
            try:
                if file_path:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        text = f.read()
                else:
                    text = file_content.decode('utf-8', errors='ignore')

                return {
                    'text': text,
                    'success': True,
                    'error': None
                }
            except Exception as e:
                logger.error(f"Error reading text file: {str(e)}")
                return {
                    'text': '',
                    'success': False,
                    'error': f"Text file reading error: {str(e)}"
                }

    except Exception as e:
        logger.error(f"Unexpected error in extract_text_from_file: {str(e)}")
        return {
            'text': '',
            'success': False,
            'error': f"Unexpected error: {str(e)}"
        }